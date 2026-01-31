# tests/test_critical_features.py
import pytest
from docx import Document
from io import BytesIO
from utils import fill_template, safe_filename, extract_text
from extraction import CVExtractor


def test_multiple_roles_same_company_creates_separate_experiences():
    """
    CRITICAL: Test that multiple roles at the same company 
    are treated as SEPARATE experience entries (not consolidated)
    """
    doc = Document()
    doc.add_paragraph("{{EXP1_COMPANY}} - {{EXP1_ROLE}}")
    doc.add_paragraph("{{EXP2_COMPANY}} - {{EXP2_ROLE}}")
    doc.add_paragraph("{{EXP3_COMPANY}} - {{EXP3_ROLE}}")
    
    # Simulate 3 different roles at Atea Pharmaceuticals
    data = {
        "candidate_name": "Jane Doe",
        "experiences": [
            {
                "company": "Atea Pharmaceuticals",
                "location": "Boston, MA",
                "role": "Manager, Quality Systems",
                "duration": "OCT 2023 - JAN 2025",
                "responsibilities": ["Task 1"]
            },
            {
                "company": "Atea Pharmaceuticals",
                "location": "Boston, MA",
                "role": "Manager, Quality System Documentation",
                "duration": "FEB 2023 - SEP 2023",
                "responsibilities": ["Task 2"]
            },
            {
                "company": "Atea Pharmaceuticals",
                "location": "Boston, MA",
                "role": "Quality Assurance Contractor",
                "duration": "OCT 2021 - FEB 2023",
                "responsibilities": ["Task 3"]
            }
        ],
        "education": [],
        "certifications": [],
        "language_skills": ["English - Fluent"]
    }
    
    result = fill_template(doc, data)
    all_text = "\n".join([p.text for p in result.paragraphs])
    
    # All 3 entries should exist separately
    assert all_text.count("Atea Pharmaceuticals") == 3
    assert "Manager, Quality Systems" in all_text
    assert "Manager, Quality System Documentation" in all_text
    assert "Quality Assurance Contractor" in all_text


def test_safe_filename_removes_special_characters():
    """Test that filenames with special characters are sanitized"""
    # Test various problematic characters
    assert safe_filename("Jane Doe: Resume.docx") == "Jane Doe_ Resume.docx"
    assert safe_filename("John/Smith\\Resume.docx") == "John_Smith_Resume.docx"
    assert safe_filename('Test"Resume".docx') == "Test_Resume_.docx"
    assert safe_filename("Resume<2024>.docx") == "Resume_2024_.docx"
    
    # Test that normal filenames pass through
    assert safe_filename("Jane_Doe_Resume.docx") == "Jane_Doe_Resume.docx"


def test_extract_text_from_docx():
    """Test that text extraction from DOCX works"""
    # Create a simple DOCX in memory
    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Software Engineer")
    doc.add_paragraph("Experience: Worked at Tech Corp")
    
    # Save to BytesIO (simulates uploaded file)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Create a mock upload object
    class MockUpload:
        def __init__(self, buffer):
            self.buffer = buffer
            self.type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            self.name = "test.docx"
        
        def getvalue(self):
            return self.buffer.getvalue()
    
    mock_file = MockUpload(buffer)
    
    # Extract text
    extracted = extract_text(mock_file)
    
    # Verify content was extracted
    assert "Jane Doe" in extracted
    assert "Software Engineer" in extracted
    assert "Experience: Worked at Tech Corp" in extracted
