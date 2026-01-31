# tests/test_edge_cases.py
import pytest
from docx import Document
from utils import fill_template
from extraction import CVExtractor


def test_missing_candidate_name_uses_placeholder():
    """
    Test that when candidate name is missing or 'Candidate Name Not Provided',
    the system handles it gracefully (filename fallback happens in cv_converter.py)
    """
    extractor = CVExtractor("fake-api-key")
    
    data = {
        "candidate_name": "Candidate Name Not Provided",
        "experiences": []
    }
    
    validated = extractor._validate_data(data)
    
    # Should preserve the placeholder (cv_converter.py will replace with filename)
    assert validated["candidate_name"] == "Candidate Name Not Provided"


def test_handles_maximum_20_experiences():
    """Test that system can handle 20 experiences (maximum supported)"""
    doc = Document()
    
    # Add placeholders for first and last experience
    doc.add_paragraph("{{EXP1_COMPANY}}")
    doc.add_paragraph("{{EXP20_COMPANY}}")
    
    # Create 20 experiences
    experiences = []
    for i in range(1, 21):
        experiences.append({
            "company": f"Company {i}",
            "location": "Boston, MA",
            "role": f"Role {i}",
            "duration": "2020 - 2023",
            "responsibilities": [f"Task {i}"]
        })
    
    data = {
        "candidate_name": "Jane Doe",
        "experiences": experiences,
        "education": [],
        "certifications": [],
        "language_skills": ["English - Fluent"]
    }
    
    result = fill_template(doc, data)
    all_text = "\n".join([p.text for p in result.paragraphs])
    
    # First and last company should both be present
    assert "Company 1" in all_text
    assert "Company 20" in all_text
    assert "{{EXP1_COMPANY}}" not in all_text
    assert "{{EXP20_COMPANY}}" not in all_text


def test_location_extracted_separately_from_company():
    """
    Test that location is stored separately from company name
    (not appended like "Formation Bio, New York, NY")
    """
    doc = Document()
    doc.add_paragraph("Company: {{EXP1_COMPANY}}")
    doc.add_paragraph("Location: {{EXP1_LOCATION}}")
    
    data = {
        "candidate_name": "Jane Doe",
        "experiences": [
            {
                "company": "Formation Bio",  # Just company name
                "location": "New York, NY",  # Separate location field
                "role": "Engineer",
                "duration": "2020 - 2023",
                "responsibilities": ["Task 1"]
            }
        ],
        "education": [],
        "certifications": [],
        "language_skills": ["English - Fluent"]
    }
    
    result = fill_template(doc, data)
    all_text = "\n".join([p.text for p in result.paragraphs])
    
    # Company and location should appear separately
    assert "Company: Formation Bio" in all_text
    assert "Location: New York, NY" in all_text
    # Should NOT be concatenated
    assert "Formation Bio, New York, NY" not in all_text


def test_handles_experience_21_beyond_max():
    """Test that experiences beyond 20 are gracefully ignored (no tokens for EXP21+)"""
    doc = Document()
    doc.add_paragraph("{{EXP1_COMPANY}}")
    doc.add_paragraph("{{EXP20_COMPANY}}")
    # Note: No {{EXP21_COMPANY}} token exists in template
    
    # Create 25 experiences (5 more than supported)
    experiences = []
    for i in range(1, 26):
        experiences.append({
            "company": f"Company {i}",
            "location": "Boston, MA",
            "role": f"Role {i}",
            "duration": "2020 - 2023",
            "responsibilities": [f"Task {i}"]
        })
    
    data = {
        "candidate_name": "Jane Doe",
        "experiences": experiences,  # 25 experiences
        "education": [],
        "certifications": [],
        "language_skills": ["English - Fluent"]
    }
    
    result = fill_template(doc, data)
    all_text = "\n".join([p.text for p in result.paragraphs])
    
    # First 20 should work
    assert "Company 1" in all_text
    assert "Company 20" in all_text
    # Beyond 20 are silently ignored (no tokens for them anyway)
    # This just ensures the function doesn't crash
