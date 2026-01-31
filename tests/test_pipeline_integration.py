# tests/test_pipeline_integration.py
import pytest
import json
from docx import Document
from utils import fill_template
from extraction import CVExtractor


def test_mock_ai_extraction_end_to_end():
    """
    Test the full pipeline: Mock AI JSON response → validate → fill template
    Simulates what happens when Gemini returns data
    """
    # Simulate a complete AI extraction response (what Gemini would return)
    mock_gemini_response = {
        "candidate_name": "JOHN DOE",  # Should be formatted to proper case
        "position": "SENIOR SOFTWARE ENGINEER",
        "total_experience_years": "8",
        "phone": "+1 555 123 4567",
        "email": "john.doe@example.com",
        "intro_paragraph": "Experienced software engineer with focus on backend systems.",
        "experiences": [
            {
                "company": "Tech Corp",
                "location": "San Francisco, CA",
                "role": "SENIOR ENGINEER",  # Should be formatted
                "duration": "Jan-2020 - Present",  # Should be normalized
                "responsibilities": [
                    "Led backend development",
                    "Managed team of 5 engineers"
                ]
            },
            {
                "company": "StartupXYZ",
                "location": "Boston, MA",
                "role": "SOFTWARE DEVELOPER",
                "duration": "Sep-2015 - Dec-2019",
                "responsibilities": [
                    "Built microservices",
                    "Implemented CI/CD pipeline"
                ]
            }
        ],
        "education": [
            {
                "institution": "MIT",
                "duration": "2011 - 2015",
                "degree": "BS Computer Science"
            }
        ],
        "technical_skills": ["Python", "Django", "PostgreSQL", "Docker"],
        "certifications": [
            {
                "name": "AWS Certified Solutions Architect",
                "year": "2022",
                "provider": "Amazon Web Services",
                "location": ""
            }
        ],
        "language_skills": ["English - Fluent", "Spanish - Conversational"]
    }
    
    # Step 1: Validate the data (what CVExtractor does)
    extractor = CVExtractor("fake-api-key")
    validated_data = extractor._validate_data(mock_gemini_response)
    
    # Verify validation happened
    assert validated_data["candidate_name"] == "John Doe"  # Formatted from JOHN DOE
    assert validated_data["position"] == "Senior Software Engineer"
    assert validated_data["experiences"][0]["role"] == "Senior Engineer"
    
    # Step 2: Create a simple template
    doc = Document()
    doc.add_paragraph("Name: {{CANDIDATE_NAME}}")
    doc.add_paragraph("Position: {{POSITION}}")
    doc.add_paragraph("Email: {{EMAIL}}")
    doc.add_paragraph("Experience 1: {{EXP1_COMPANY}} - {{EXP1_ROLE}}")
    doc.add_paragraph("Experience 2: {{EXP2_COMPANY}} - {{EXP2_ROLE}}")
    doc.add_paragraph("Education: {{EDU1_INSTITUTION}} - {{EDU1_DEGREE}}")
    doc.add_paragraph("Cert: {{CERT1_NAME}}")
    
    # Step 3: Fill template with validated data
    result = fill_template(doc, validated_data)
    
    # Step 4: Verify final output
    all_text = "\n".join([p.text for p in result.paragraphs])
    
    # Check that all data made it through the pipeline
    assert "John Doe" in all_text
    assert "Senior Software Engineer" in all_text
    assert "john.doe@example.com" in all_text
    assert "Tech Corp" in all_text
    assert "Senior Engineer" in all_text
    assert "StartupXYZ" in all_text
    assert "MIT" in all_text
    assert "BS Computer Science" in all_text
    assert "AWS Certified Solutions Architect" in all_text
    
    # Critical: No tokens should remain
    assert "{{" not in all_text


def test_empty_certifications_and_education_handled_gracefully():
    """
    Test that CVs with no certifications or education don't break the system
    Common in some industries or career switchers
    """
    doc = Document()
    doc.add_paragraph("Name: {{CANDIDATE_NAME}}")
    doc.add_paragraph("Cert 1: {{CERT1_NAME}}")
    doc.add_paragraph("Cert 2: {{CERT2_NAME}}")
    doc.add_paragraph("Education: {{EDU1_INSTITUTION}}")
    doc.add_paragraph("Education: {{EDU2_INSTITUTION}}")
    
    # Candidate with NO certifications and NO formal education
    data = {
        "candidate_name": "Jane Doe",
        "position": "Self-taught Developer",
        "experiences": [
            {
                "company": "Freelance",
                "location": "Remote",
                "role": "Full Stack Developer",
                "duration": "2020 - Present",
                "responsibilities": ["Built 50+ websites"]
            }
        ],
        "education": [],  # Empty
        "certifications": [],  # Empty
        "technical_skills": ["JavaScript", "React", "Node.js"],
        "language_skills": ["English - Fluent"]
    }
    
    result = fill_template(doc, data)
    all_text = "\n".join([p.text for p in result.paragraphs])
    
    # Name should be present
    assert "Jane Doe" in all_text
    
    # Empty cert/education lines should be removed or handled
    # The key is that NO unreplaced tokens remain
    assert "{{CERT1_NAME}}" not in all_text
    assert "{{CERT2_NAME}}" not in all_text
    assert "{{EDU1_INSTITUTION}}" not in all_text
    assert "{{EDU2_INSTITUTION}}" not in all_text
    
    # No tokens at all should remain
    assert "{{" not in all_text
