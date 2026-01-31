# tests/test_template.py
import pytest
from docx import Document
from utils import fill_template


def test_fill_template_replaces_basic_tokens():
    """Test that basic scalar tokens get replaced"""
    doc = Document()
    doc.add_paragraph("Name: {{CANDIDATE_NAME}}")
    doc.add_paragraph("Position: {{POSITION}}")
    doc.add_paragraph("Email: {{EMAIL}}")
    
    data = {
        "candidate_name": "Jane Doe",
        "position": "Software Engineer",
        "email": "jane@example.com",
        "experiences": [],
        "education": [],
        "certifications": [],
        "language_skills": ["English - Fluent"]
    }
    
    result = fill_template(doc, data)
    
    # Extract all text from document
    all_text = "\n".join([p.text for p in result.paragraphs])
    
    assert "Jane Doe" in all_text
    assert "Software Engineer" in all_text
    assert "jane@example.com" in all_text
    assert "{{CANDIDATE_NAME}}" not in all_text


def test_fill_template_replaces_experience_tokens():
    """Test that experience tokens get replaced correctly"""
    doc = Document()
    doc.add_paragraph("{{EXP1_COMPANY}}")
    doc.add_paragraph("{{EXP1_ROLE}}")
    doc.add_paragraph("{{EXP1_DURATION}}")
    doc.add_paragraph("{{EXP1_LOCATION}}")
    doc.add_paragraph("{{EXP1_RESP1}}")
    
    data = {
        "candidate_name": "Jane Doe",
        "experiences": [
            {
                "company": "Formation Bio",
                "location": "New York, NY",
                "role": "QA Engineer",
                "duration": "SEP 2020 - Present",
                "responsibilities": ["Did task A"]
            }
        ],
        "education": [],
        "certifications": [],
        "language_skills": ["English - Fluent"]
    }
    
    result = fill_template(doc, data)
    all_text = "\n".join([p.text for p in result.paragraphs])
    
    assert "Formation Bio" in all_text
    assert "QA Engineer" in all_text
    assert "SEP 2020 - Present" in all_text
    assert "New York, NY" in all_text
    assert "Did task A" in all_text


def test_fill_template_removes_empty_responsibility_lines():
    """Test that empty responsibility bullets get removed"""
    doc = Document()
    doc.add_paragraph("{{EXP1_COMPANY}}")
    doc.add_paragraph("{{EXP1_RESP1}}")
    doc.add_paragraph("{{EXP1_RESP2}}")
    doc.add_paragraph("{{EXP1_RESP3}}")  # This should be removed (no data)
    
    data = {
        "candidate_name": "Jane Doe",
        "experiences": [
            {
                "company": "Tech Corp",
                "location": "Boston, MA",
                "role": "Engineer",
                "duration": "2020 - 2023",
                "responsibilities": ["Task 1", "Task 2"]  # Only 2 responsibilities
            }
        ],
        "education": [],
        "certifications": [],
        "language_skills": ["English - Fluent"]
    }
    
    result = fill_template(doc, data)
    all_text = "\n".join([p.text for p in result.paragraphs])
    
    assert "Task 1" in all_text
    assert "Task 2" in all_text
    # RESP3 should be removed since there's no 3rd responsibility
    assert "{{EXP1_RESP3}}" not in all_text


def test_fill_template_handles_education():
    """Test that education tokens get replaced"""
    doc = Document()
    doc.add_paragraph("{{EDU1_INSTITUTION}}")
    doc.add_paragraph("{{EDU1_DEGREE}}")
    doc.add_paragraph("{{EDU1_DURATION}}")
    
    data = {
        "candidate_name": "Jane Doe",
        "experiences": [],
        "education": [
            {
                "institution": "MIT",
                "degree": "BS Computer Science",
                "duration": "2015 - 2019"
            }
        ],
        "certifications": [],
        "language_skills": ["English - Fluent"]
    }
    
    result = fill_template(doc, data)
    all_text = "\n".join([p.text for p in result.paragraphs])
    
    assert "MIT" in all_text
    assert "BS Computer Science" in all_text
    assert "2015 - 2019" in all_text


def test_fill_template_handles_certifications():
    """Test that certification tokens get replaced"""
    doc = Document()
    doc.add_paragraph("{{CERT1_NAME}}")
    doc.add_paragraph("{{CERT1_PROVIDER}}")
    doc.add_paragraph("{{CERT1_YEAR}}")
    
    data = {
        "candidate_name": "Jane Doe",
        "experiences": [],
        "education": [],
        "certifications": [
            {
                "name": "PMP Certification",
                "provider": "PMI",
                "year": "2023",
                "location": ""
            }
        ],
        "language_skills": ["English - Fluent"]
    }
    
    result = fill_template(doc, data)
    all_text = "\n".join([p.text for p in result.paragraphs])
    
    assert "PMP Certification" in all_text
    assert "PMI" in all_text
    assert "2023" in all_text


def test_fill_template_no_unreplaced_tokens():
    """Critical test: ensure NO tokens remain unreplaced"""
    doc = Document()
    doc.add_paragraph("{{CANDIDATE_NAME}}")
    doc.add_paragraph("{{POSITION}}")
    doc.add_paragraph("{{EXP1_COMPANY}}")
    
    data = {
        "candidate_name": "Jane Doe",
        "position": "Engineer",
        "experiences": [
            {
                "company": "Tech Corp",
                "location": "Boston, MA",
                "role": "Engineer",
                "duration": "2020 - 2023",
                "responsibilities": ["Task 1"]
            }
        ],
        "education": [],
        "certifications": [],
        "language_skills": ["English - Fluent"]
    }
    
    result = fill_template(doc, data)
    all_text = "\n".join([p.text for p in result.paragraphs])
    
    # This is the most important check - no {{tokens}} should remain
    assert "{{" not in all_text
    assert "}}" not in all_text
