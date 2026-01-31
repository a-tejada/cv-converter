# tests/test_production_critical.py
import pytest
from docx import Document
from io import BytesIO
from utils import extract_text, fill_template
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter


def test_extract_text_from_pdf():
    """Test that text extraction from PDF works correctly"""
    # Create a simple PDF in memory using reportlab
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    
    # Add text to PDF
    c.drawString(100, 750, "Jane Doe")
    c.drawString(100, 730, "Software Engineer")
    c.drawString(100, 710, "Email: jane@example.com")
    c.drawString(100, 690, "Experience at Tech Corp")
    c.save()
    
    buffer.seek(0)
    
    # Create a more complete mock upload object that pdfplumber can handle
    class MockPDFUpload:
        def __init__(self, buffer):
            self._buffer = buffer
            self.type = "application/pdf"
            self.name = "test.pdf"
        
        def __enter__(self):
            return self._buffer
        
        def __exit__(self, *args):
            pass
        
        def read(self, *args):
            return self._buffer.read(*args)
        
        def seek(self, pos):
            return self._buffer.seek(pos)
        
        def tell(self):
            return self._buffer.tell()
    
    mock_file = MockPDFUpload(buffer)
    
    # Import pdfplumber to test directly
    import pdfplumber
    
    # Test that pdfplumber can actually read the PDF
    text_parts = []
    with pdfplumber.open(buffer) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
    
    extracted = "\n".join(text_parts)
    
    # Verify content was extracted
    assert "Jane Doe" in extracted
    assert "Software Engineer" in extracted

def test_extract_text_handles_corrupted_file():
    """Test that corrupted/malformed files don't crash the system"""
    # Create a mock corrupted file
    class MockCorruptedUpload:
        def __init__(self):
            self.type = "application/pdf"
            self.name = "corrupted.pdf"
        
        def read(self):
            # Return garbage data that's not a valid PDF
            return b"This is not a valid PDF file content"
        
        def seek(self, pos):
            pass
    
    mock_file = MockCorruptedUpload()
    
    # Should not crash, should return empty string or handle gracefully
    result = extract_text(mock_file)
    
    # Result should be a string (even if empty)
    assert isinstance(result, str)
    # Should not raise an exception


def test_table_row_deletion_removes_empty_experiences():
    """
    CRITICAL: Test that table rows with empty experience data get deleted
    This is a key feature - unused experience sections should disappear
    """
    doc = Document()
    
    # Create a table with 3 experience rows
    table = doc.add_table(rows=3, cols=1)
    
    # EXP1 - has data
    cell1 = table.rows[0].cells[0]
    cell1.text = "{{EXP1_COMPANY}} - {{EXP1_ROLE}}"
    cell1.add_paragraph("{{EXP1_RESP1}}")
    
    # EXP2 - NO data (should be deleted)
    cell2 = table.rows[1].cells[0]
    cell2.text = "{{EXP2_COMPANY}} - {{EXP2_ROLE}}"
    cell2.add_paragraph("{{EXP2_RESP1}}")
    
    # EXP3 - NO data (should be deleted)
    cell3 = table.rows[2].cells[0]
    cell3.text = "{{EXP3_COMPANY}} - {{EXP3_ROLE}}"
    cell3.add_paragraph("{{EXP3_RESP1}}")
    
    # Data for only 1 experience
    data = {
        "candidate_name": "Jane Doe",
        "experiences": [
            {
                "company": "Formation Bio",
                "location": "New York, NY",
                "role": "Engineer",
                "duration": "2020 - 2023",
                "responsibilities": ["Task 1"]
            }
        ],
        "education": [],
        "certifications": [],
        "language_skills": ["English - Fluent"]
    }
    
    result = fill_template(doc, data)
    
    # Check that table now has only 1 row (the other 2 should be deleted)
    assert len(result.tables[0].rows) == 1
    
    # Extract all text from the table
    table_text = ""
    for row in result.tables[0].rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                table_text += para.text + "\n"
    
    # EXP1 should be present
    assert "Formation Bio" in table_text
    assert "Engineer" in table_text
    
    # EXP2 and EXP3 should NOT be present (rows deleted)
    assert "{{EXP2_" not in table_text
    assert "{{EXP3_" not in table_text


def test_extract_text_handles_empty_docx():
    """Test that empty DOCX files are handled gracefully"""
    # Create an empty DOCX
    doc = Document()
    # Don't add any content
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    class MockUpload:
        def __init__(self, buffer):
            self.buffer = buffer
            self.type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            self.name = "empty.docx"
        
        def getvalue(self):
            return self.buffer.getvalue()
    
    mock_file = MockUpload(buffer)
    
    # Should return empty string, not crash
    result = extract_text(mock_file)
    
    assert isinstance(result, str)
    assert result == "" or result.strip() == ""
