# utils.py
# Helper functions for CV processing and template filling

import re
from io import BytesIO
from typing import Dict, Any
import streamlit as st
import pdfplumber
from docx import Document
from docx.shared import Pt, RGBColor

# ────────────────────────────────────────────────────────────────
#  Original Helper Functions
# ────────────────────────────────────────────────────────────────
def mask_api_key(api_key: str) -> str:
    """Mask API key for display, showing only first 4 and last 4 characters."""
    if not api_key or len(api_key) < 12:
        return api_key
    return f"{api_key[:4]}{'*' * (len(api_key) - 8)}{api_key[-4:]}"


def extract_text(upload) -> str:
    try:
        if upload.type == "application/pdf":
            text_parts = []
            with pdfplumber.open(upload) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        # Clean the text
                        text = text.encode('utf-8', errors='ignore').decode('utf-8')
                        text_parts.append(text)
            return "\n".join(text_parts)
            
        if upload.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(BytesIO(upload.getvalue()))
            return "\n".join(p.text for p in doc.paragraphs)
            
        return upload.read().decode("utf-8", errors="ignore")
    except Exception as e:
        st.error(f"Error reading {upload.name}: {e}")
        return ""
    
def format_date(date_str: str) -> str:
    """Convert various date formats to MMM YYYY format."""
    if not date_str:
        return ""
    
    # Check for present/current/ongoing terms
    if date_str.lower() in ["present", "current", "ongoing", "till date", "now", "till now"]:
        return "Present"
    
    # Common date patterns
    patterns = [
        # Handle Sep-2015, Sep 2015, Sep/2015 formats
        (r'(\w{3,})[- /](\d{4})', lambda m: f"{m.group(1).upper()[:3]} {m.group(2)}"),
        # Handle 09/2015, 09-2015 formats  
        (r'(\d{1,2})[/-](\d{4})', lambda m: f"{get_month_abbr(m.group(1).zfill(2))} {m.group(2)}"),
        # Handle September 2015, Sep 2015 formats
        (r'(\w+)\s+(\d{4})', lambda m: f"{m.group(1)[:3].upper()} {m.group(2)}"),
        # Handle September, 2015 formats
        (r'(\w+),?\s+(\d{4})', lambda m: f"{m.group(1)[:3].upper()} {m.group(2)}"),
    ]
    
    for pattern, formatter in patterns:
        match = re.search(pattern, date_str, re.IGNORECASE)
        if match:
            return formatter(match)
    
    return date_str  # Return as-is if no pattern matches

def get_month_abbr(month_num: str) -> str:
    """Convert month number to 3-letter abbreviation."""
    months = ["JAN", "FEB", "MAR", "APR", "MAY", "JUN", 
              "JUL", "AUG", "SEP", "OCT", "NOV", "DEC"]
    try:
        # Handle both "02" and "2" formats
        month_int = int(month_num.lstrip('0') or '0')
        if 1 <= month_int <= 12:
            return months[month_int - 1]
    except:
        pass
    return month_num

def format_duration(duration: str, is_first_experience: bool = False) -> str:
    """Format duration string to MMM YYYY - MMM YYYY format."""
    if not duration:
        return ""
    
    # Preserve "Present" for ongoing positions
    if " - Present" in duration or "- Present" in duration:
        # Just format the start date part
        parts = re.split(r'\s*-\s*', duration)
        if parts:
            start = parts[0].strip()
            # Convert month to uppercase
            start = re.sub(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)', 
                          lambda m: m.group(1).upper(), start, flags=re.IGNORECASE)
            # Ensure space between month and year
            start = re.sub(r'([A-Z]{3})-?(\d{4})', r'\1 \2', start)
            return f"{start} to Present"
    
    # Split by common separators
    parts = re.split(r'\s*[-–—]\s*', duration)
    
    if len(parts) == 2:
        start = format_date(parts[0].strip())
        end = format_date(parts[1].strip())
        return f"{start} to {end}"
    elif len(parts) == 1 and is_first_experience:
        # For first experience, if only start date, add Present
        start = format_date(parts[0].strip())
        if start and start != "Present":
            return f"{start} to Present"
        return start
    
    return duration

def format_name(name: str) -> str:
    """Convert name from ALL CAPS to Proper Case."""
    if not name:
        return ""
    
    # Handle common name patterns
    words = name.split()
    formatted_words = []
    
    for word in words:
        if word.isupper() and len(word) > 1:
            # Convert from ALL CAPS to Proper Case
            formatted_words.append(word.capitalize())
        else:
            formatted_words.append(word)
    
    return " ".join(formatted_words)


# ────────────────────────────────────────────────────────────────
#  Helper: Check if a table row contains experience placeholders
# ────────────────────────────────────────────────────────────────
def contains_experience_placeholder(text: str, exp_num: int) -> bool:
    """Check if text contains placeholders for a specific experience number."""
    patterns = [
        f"{{{{EXP{exp_num}_COMPANY}}}}",
        f"{{{{EXP{exp_num}_ROLE}}}}",
        f"{{{{EXP{exp_num}_DURATION}}}}",
        f"{{{{EXP{exp_num}_RESP"
    ]
    return any(pattern in text for pattern in patterns)

def get_row_text(row) -> str:
    """Get all text from a table row."""
    text = ""
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            text += paragraph.text + " "
    return text

def should_delete_row(row, exp_num: int, has_data: bool) -> bool:
    """Determine if a row should be deleted based on experience data availability."""
    row_text = get_row_text(row)
    
    # Check if this row contains placeholders for this experience number
    if contains_experience_placeholder(row_text, exp_num):
        # If no data for this experience, mark for deletion
        return not has_data
    
    return False

# ────────────────────────────────────────────────────────────────
#  Enhanced template filling with row deletion
# ────────────────────────────────────────────────────────────────
def set_paragraph_format(paragraph, text, font_name="Arial", font_size=10, bold=False):
    """Set consistent formatting for a paragraph."""
    paragraph.text = text
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)  # Black

def fill_template(doc: Document, d: Dict[str, Any]) -> Document:
    """Fill template with proper formatting and delete unused experience rows."""
    
    # Basic replacements
    basic_repl = {
        "{{CANDIDATE_NAME}}": str(d.get("candidate_name", "") or ""),
        "{{POSITION}}": str(d.get("position", "") or ""),
        "{{TOTAL_EXPERIENCE_YEARS}}": str(d.get("total_experience_years", "") or ""),
        "{{PHONE}}": str(d.get("phone", "") or ""),
        "{{EMAIL}}": str(d.get("email", "") or ""),
        "{{INTRO_PARAGRAPH}}": str(d.get("intro_paragraph", "") or ""),
    }
    
    # Track which experiences have data
    experiences_with_data = set()
    
    # Process experiences (up to 20)
    exp_repl = {}
    for i in range(1, 21):  # 1 to 20
        if i <= len(d.get("experiences", [])):
            exp = d["experiences"][i-1]
            # Check if this experience has meaningful data
            if exp.get("company") and exp.get("role"):
                experiences_with_data.add(i)
                
                # Mark company placeholders for bold formatting
                company_text = str(exp.get("company", "") or "")
                exp_repl[f"{{{{EXP{i}_COMPANY}}}}"] = f"<<<BOLD>>>{company_text}<<<END_BOLD>>>"
                exp_repl[f"{{{{EXP{i}_ROLE}}}}"] = str(exp.get("role", "") or "")
                exp_repl[f"{{{{EXP{i}_DURATION}}}}"] = str(exp.get("duration", "") or "")
                # Fallback for location if not specified
                location = str(exp.get("location", "") or "")
                exp_repl[f"{{{{EXP{i}_LOCATION}}}}"] = location if location.strip() else "Location Not Specified"
                
                # Handle responsibilities
                responsibilities = exp.get("responsibilities", [])
                
                # Handle all 100 responsibility placeholders
                for j in range(1, 101):  # 1 to 100
                    placeholder = f"{{{{EXP{i}_RESP{j}}}}}"
                    if j <= len(responsibilities):
                        resp_text = responsibilities[j-1]
                        if resp_text is None:
                            resp_text = ""
                        exp_repl[placeholder] = str(resp_text)
                    else:
                        # Mark empty responsibilities for removal
                        exp_repl[placeholder] = "<<<REMOVE_THIS_LINE>>>"
            else:
                # Mark all placeholders for this experience for deletion
                exp_repl[f"{{{{EXP{i}_COMPANY}}}}"] = "<<<DELETE_EXPERIENCE>>>"
                exp_repl[f"{{{{EXP{i}_ROLE}}}}"] = "<<<DELETE_EXPERIENCE>>>"
                exp_repl[f"{{{{EXP{i}_DURATION}}}}"] = "<<<DELETE_EXPERIENCE>>>"
                exp_repl[f"{{{{EXP{i}_LOCATION}}}}"] = "<<<DELETE_EXPERIENCE>>>"  # ADD THIS LINE
                for j in range(1, 101):
                    exp_repl[f"{{{{EXP{i}_RESP{j}}}}}"] = "<<<DELETE_EXPERIENCE>>>"
        else:
            # No data for this experience - mark for deletion
            exp_repl[f"{{{{EXP{i}_COMPANY}}}}"] = "<<<DELETE_EXPERIENCE>>>"
            exp_repl[f"{{{{EXP{i}_ROLE}}}}"] = "<<<DELETE_EXPERIENCE>>>"
            exp_repl[f"{{{{EXP{i}_DURATION}}}}"] = "<<<DELETE_EXPERIENCE>>>"
            exp_repl[f"{{{{EXP{i}_LOCATION}}}}"] = "<<<DELETE_EXPERIENCE>>>"  # ADD THIS LINE
            for j in range(1, 101):
                exp_repl[f"{{{{EXP{i}_RESP{j}}}}}"] = "<<<DELETE_EXPERIENCE>>>"

    # Process education (up to 5 degrees)
    edu_repl = {}
    education_data = d.get("education", [])
    if not isinstance(education_data, list):
        education_data = []

    for i in range(1, 6):  # 1 to 5
        if i <= len(education_data):
            edu = education_data[i-1]
            edu_repl[f"{{{{EDU{i}_INSTITUTION}}}}"] = str(edu.get("institution", "") or "")
            # Fallback if duration is empty
            duration = str(edu.get("duration", "") or "")
            edu_repl[f"{{{{EDU{i}_DURATION}}}}"] = duration if duration.strip() else "Dates Not Available"
            edu_repl[f"{{{{EDU{i}_DEGREE}}}}"] = str(edu.get("degree", "") or "")

        else:
            # Mark empty education slots for removal
            edu_repl[f"{{{{EDU{i}_INSTITUTION}}}}"] = "<<<REMOVE_THIS_LINE>>>"
            edu_repl[f"{{{{EDU{i}_DURATION}}}}"] = "<<<REMOVE_THIS_LINE>>>"
            edu_repl[f"{{{{EDU{i}_DEGREE}}}}"] = "<<<REMOVE_THIS_LINE>>>"

    # Process certifications (up to 10)
    cert_repl = {}
    certifications_data = d.get("certifications", [])
    if not isinstance(certifications_data, list):
        certifications_data = []

    for i in range(1, 11):  # 1 to 10
        if i <= len(certifications_data):
            cert = certifications_data[i-1]
            cert_repl[f"{{{{CERT{i}_NAME}}}}"] = str(cert.get("name", "") or "")
            # Fallback for year if not specified
            year = str(cert.get("year", "") or "")
            cert_repl[f"{{{{CERT{i}_YEAR}}}}"] = year if year.strip() else "Year Not Available"
            # Fallback for provider if not specified
            provider = str(cert.get("provider", "") or "")
            cert_repl[f"{{{{CERT{i}_PROVIDER}}}}"] = provider if provider.strip() else "Provider Not Specified"
            cert_repl[f"{{{{CERT{i}_LOCATION}}}}"] = str(cert.get("location", "") or "")
        else:
            # Mark empty certification slots for removal
            cert_repl[f"{{{{CERT{i}_NAME}}}}"] = "<<<REMOVE_THIS_LINE>>>"
            cert_repl[f"{{{{CERT{i}_YEAR}}}}"] = "<<<REMOVE_THIS_LINE>>>"
            cert_repl[f"{{{{CERT{i}_PROVIDER}}}}"] = "<<<REMOVE_THIS_LINE>>>"
            cert_repl[f"{{{{CERT{i}_LOCATION}}}}"] = "<<<REMOVE_THIS_LINE>>>"


    # Format skills for backward compatibility (if old template tokens still exist)
    tech_skills = d.get("technical_skills", [])
    if tech_skills:
        tech_skills_text = "\n".join([f"• {skill}" for skill in tech_skills])
    else:
        tech_skills_text = ""

    langs = d.get("language_skills", [])
    if langs:
        langs_text = ", ".join(langs)
    else:
        langs_text = "English - Fluent"

    # Add formatted lists to replacements
    basic_repl["{{TECHNICAL_SKILLS_LIST}}"] = tech_skills_text
    basic_repl["{{LANGUAGE_SKILLS_LIST}}"] = langs_text

    # Combine all replacements
    all_repl = {**basic_repl, **exp_repl, **edu_repl, **cert_repl}

    
    # Process paragraphs
    paragraphs_to_remove = []
    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        new_text = original_text
        
        # Apply replacements
        for placeholder, value in all_repl.items():
            if placeholder in new_text:
                if value is None:
                    value = ""
                new_text = new_text.replace(placeholder, str(value))
        
        # Check if this paragraph is part of a deleted experience section
        if "<<<DELETE_EXPERIENCE>>>" in new_text:
            paragraphs_to_remove.append(paragraph)
            continue
        
        # Check if this line should be removed (empty responsibility)
        if "<<<REMOVE_THIS_LINE>>>" in new_text:
            new_text = new_text.replace("<<<REMOVE_THIS_LINE>>>", "")
            # If the resulting text is empty or just whitespace/bullet, remove the paragraph
            if not new_text.strip() or new_text.strip() in ["-", "•"]:
                paragraphs_to_remove.append(paragraph)
                continue
        
        # If text changed, update with formatting
        if new_text != original_text:
            paragraph.clear()
            # Remove the bullet point if the line only contains "-"
            if new_text.strip() != "-":
                # Check for bold markers
                if "<<<BOLD>>>" in new_text and "<<<END_BOLD>>>" in new_text:
                    # Extract and apply bold formatting
                    parts = new_text.split("<<<BOLD>>>")
                    for i, part in enumerate(parts):
                        if i == 0 and part:
                            # Text before first bold marker
                            run = paragraph.add_run(part)
                            run.font.name = 'Arial'
                            run.font.size = Pt(10)
                            run.font.color.rgb = RGBColor(0, 0, 0)
                        elif "<<<END_BOLD>>>" in part:
                            # This part contains bold text
                            bold_parts = part.split("<<<END_BOLD>>>")
                            # Add bold text
                            run = paragraph.add_run(bold_parts[0])
                            run.font.name = 'Arial'
                            run.font.size = Pt(10)
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            # Add remaining text if any
                            if len(bold_parts) > 1 and bold_parts[1]:
                                run = paragraph.add_run(bold_parts[1])
                                run.font.name = 'Arial'
                                run.font.size = Pt(10)
                                run.font.color.rgb = RGBColor(0, 0, 0)
                else:
                    # No bold markers, normal text
                    paragraph.add_run(new_text)
                    # Apply Arial 10pt black formatting
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)
                        run.font.bold = False
                        run.font.color.rgb = RGBColor(0, 0, 0)
                # Set 1.5 line spacing
                paragraph.paragraph_format.line_spacing = 1.5

    # Remove empty paragraphs
    for p in paragraphs_to_remove:
        p._element.getparent().remove(p._element)
    
    # Process tables and handle row deletion
    for table in doc.tables:
        rows_to_delete = []
        
        # First pass: identify rows to delete
        for row_idx, row in enumerate(table.rows):
            row_text = get_row_text(row)
            
            # Check each experience number
            for exp_num in range(1, 21):
                if contains_experience_placeholder(row_text, exp_num):
                    if exp_num not in experiences_with_data:
                        # This row contains placeholders for an experience we don't have data for
                        rows_to_delete.append(row_idx)
                        break
        
        # Second pass: process cells in rows we're keeping
        for row_idx, row in enumerate(table.rows):
            if row_idx in rows_to_delete:
                continue  # Skip rows marked for deletion
                
            for cell in row.cells:
                paragraphs_to_remove = []
                for paragraph in cell.paragraphs:
                    original_text = paragraph.text
                    new_text = original_text
                    
                    # Apply replacements
                    for placeholder, value in all_repl.items():
                        if placeholder in new_text:
                            new_text = new_text.replace(placeholder, value)
                    
                    # Skip if it's marked for deletion
                    if "<<<DELETE_EXPERIENCE>>>" in new_text:
                        continue
                    
                    # Check if this line should be removed
                    if "<<<REMOVE_THIS_LINE>>>" in new_text:
                        new_text = new_text.replace("<<<REMOVE_THIS_LINE>>>", "")
                        # If the resulting text is empty or just whitespace/bullet, remove the paragraph
                        if not new_text.strip() or new_text.strip() in ["-", "•"]:
                            paragraphs_to_remove.append(paragraph)
                            continue
    
                    # If text changed, update with formatting
                    if new_text != original_text:
                        paragraph.clear()
                        
                        # Skip if it's just a bullet point with no content
                        if new_text.strip() == "-":
                            paragraphs_to_remove.append(paragraph)
                            continue
                        
                        # Check for bold markers in table cells
                        if "<<<BOLD>>>" in new_text and "<<<END_BOLD>>>" in new_text:
                            # Extract and apply bold formatting
                            parts = new_text.split("<<<BOLD>>>")
                            for i, part in enumerate(parts):
                                if i == 0 and part:
                                    # Text before first bold marker
                                    run = paragraph.add_run(part)
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(10)
                                    run.font.color.rgb = RGBColor(0, 0, 0)
                                elif "<<<END_BOLD>>>" in part:
                                    # This part contains bold text
                                    bold_parts = part.split("<<<END_BOLD>>>")
                                    # Add bold text
                                    run = paragraph.add_run(bold_parts[0])
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(10)
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(0, 0, 0)
                                    # Add remaining text if any
                                    if len(bold_parts) > 1 and bold_parts[1]:
                                        run = paragraph.add_run(bold_parts[1])
                                        run.font.name = 'Arial'
                                        run.font.size = Pt(10)
                                        run.font.color.rgb = RGBColor(0, 0, 0)
                        # Handle multi-line content (like skills/certs)
                        elif '\n' in new_text:
                            lines = new_text.split('\n')
                            for idx, line in enumerate(lines):
                                if idx > 0:
                                    paragraph = cell.add_paragraph()
                                
                                # Check if this is a project header line with Duration
                                if (line.strip().startswith("Project Name:") and 
                                    ("Location:" in line.strip() and "Duration:" in line.strip())):
                                    # Make the entire project header bold (including Duration)
                                    run = paragraph.add_run(line)
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(10)
                                    run.font.bold = True  # Make entire line bold
                                    run.font.color.rgb = RGBColor(0, 0, 0)
                                elif line.strip().startswith("Project Name:") and "Location:" in line.strip():
                                    # Old format without Duration - still make bold
                                    run = paragraph.add_run(line)
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(10)
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(0, 0, 0)
                                else:
                                    run = paragraph.add_run(line)
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(10)
                                    run.font.bold = False
                                    run.font.color.rgb = RGBColor(0, 0, 0)
                                # Set 1.5 line spacing
                                paragraph.paragraph_format.line_spacing = 1.5
                        else:
                            # Single line handling
                            if (new_text.strip().startswith("Project Name:") and 
                                ("Location:" in new_text.strip() and "Duration:" in new_text.strip())):
                                # Project header with Duration - make entire line bold
                                if new_text.strip().startswith("• "):
                                    new_text = new_text.replace("• ", "", 1)
                                run = paragraph.add_run(new_text)
                                run.font.name = 'Arial'
                                run.font.size = Pt(10)
                                run.font.bold = True  # Make entire line bold
                                run.font.color.rgb = RGBColor(0, 0, 0)
                            elif new_text.strip().startswith("Project Name:") and "Location:" in new_text.strip():
                                # Old format without Duration - still make bold
                                if new_text.strip().startswith("• "):
                                    new_text = new_text.replace("• ", "", 1)
                                run = paragraph.add_run(new_text)
                                run.font.name = 'Arial'
                                run.font.size = Pt(10)
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(0, 0, 0)
                            else:
                                run = paragraph.add_run(new_text)
                                run.font.name = 'Arial'
                                run.font.size = Pt(10)
                                run.font.bold = False
                                run.font.color.rgb = RGBColor(0, 0, 0)
                            # Set 1.5 line spacing
                            paragraph.paragraph_format.line_spacing = 1.5                
                # Remove empty paragraphs from cells
                for p in paragraphs_to_remove:
                    try:
                        p._element.getparent().remove(p._element)
                    except:
                        pass  # Some paragraphs might be required by the table structure
        
        # Third pass: Actually delete the rows (in reverse order to maintain indices)
        for row_idx in sorted(rows_to_delete, reverse=True):
            try:
                row = table.rows[row_idx]
                tbl = table._tbl
                tbl.remove(row._tr)
            except Exception as e:
                st.warning(f"Could not delete row {row_idx}: {str(e)}")
    
    return doc

def safe_filename(name: str) -> str:
    return re.sub(r'[\\/*?:"<>|]', "_", name).strip() or "output"
